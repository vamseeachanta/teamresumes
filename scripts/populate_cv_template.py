"""
Script to populate Project Pheasant CV Template with Vamsee Achanta's resume data
"""
from docx import Document
from datetime import datetime

def populate_cv_template():
    """Populate the CV template with resume data"""
    
    # Load the template
    doc = Document('cv/acma/01_format/Project Pheasant CV Template.docx')
    
    # Resume data
    cv_data = {
        'name': 'Vamsee Achanta, P.E.',
        'profession': 'Naval Architect / Subsea Systems Engineering Expert',
        'dob': '',  # Not provided in resume
        'nationality': 'US Citizen',
        'memberships': 'Society of Petroleum Engineers (SPE), American Society of Mechanical Engineers (ASME), Marine Technology Society (MTS), API Standards Committees (API-RP-16Q, API-RP-17G, API-RP-17G2)',
        'function': 'Naval Architect / Marine Systems Analysis Lead',
        'key_qualifications': """VP of Engineering with 23+ years driving subsea/offshore projects from concept to decommission. Led teams of 20+ engineers delivering 100% on-time project completion across 100+ riser design projects worth $500M+ in total value.

Key Areas of Expertise:
• Marine Dynamic Analysis - Coupled analysis of floating systems (FSTs, FPSOs, LNG carriers) with moorings in extreme weather
• Subsea Systems - Risers (SCR, TTR, Hybrid), Pipelines, Umbilicals, Manifolds, Jumpers
• Digital Transformation - Implemented 50+ physics-based algorithms for real-time drilling and production analytics (20,000 wells)
• Crisis Management - Engineering Manager for BP Macondo containment riser emergency response
• Software Proficiency - AQWA, OrcaWave, OrcaFlex, ANSYS, Abaqus, COMSOL, Python, SQL
• API Standards - Committee member for API-RP-16Q, API-RP-17G, API-RP-17G2

Recent Project Highlights:
• WoodFibre LNG Terminal ($1.8B) - Marine global analysis for dual FST berthing system handling 180,000m³ LNG carriers (99.5% availability)
• Digital twins development using Python/AI reducing analysis time by 70%
• Floundering analysis for litigation assistance using advanced time-domain marine analysis""",
        
        'education': """Master of Science in Mechanical Engineering
Texas A&M University, College Station, TX
Specialization: Offshore Structures & Marine Systems

Bachelor of Technology in Mechanical Engineering
Indian Institute of Technology (IIT), Madras, India

Professional Credentials:
• Professional Engineer (P.E.) - Texas
• API Committee Member - API-RP-16Q, API-RP-17G, API-RP-17G2""",
        
        'employment': """Dec 2023 - Present: Naval Architect, Alan McClure & Associates, Houston, TX
• WoodFibre LNG Terminal Design ($1.8B project, British Columbia) - Detailed marine global analysis for 2 FSTs berthed side-by-side in 35m water depth with coupled analysis for extreme weather conditions
• Designed mooring system handling 180,000m³ LNG carriers with 99.5% availability
• Diffraction and mooring analysis for barges investigating mooring line failures in extreme weather
• Floundering analysis of tug boats due to passing ships using time-domain analysis
• Digital transformation: Developed digital twins for hull diffraction analysis in AQWA/OrcaWave and mooring analysis using OrcFxAPI and Python/AI
• Software: AQWA, OrcaWave, OrcaFlex
Client References: Available upon request

Jun 2016 - Present: VP of Engineering, FDAS, Houston, TX
• 6000 ft Water Depth Semisubmersible Production Vessel Design ($150M+ project) - Led feasibility study saving $30M through conversion strategy
• Economic Analysis Platform for GoM Fields - Built SQL/Python analytics platform analyzing 200+ fields, reducing evaluation time from 2 weeks to 2 days
• SEWOL Salvage Ship Critical Analysis ($20M emergency) - Delivered critical hull FEA analysis in 72 hours enabling safe lifting of 6,800-ton vessel
• Secured $2M in follow-on contracts through technical excellence
Client References: Korean government salvage operation, GoM operators

Jun 2012 - Present: Engineering Lead Consultant, AceEngineer, Houston, TX
• ExxonMobil Yellowtail ($2B field development, Guyana) - Analyzed 6 umbilicals in 6,000ft water depth, saving 12 vessel days ($6M savings) with zero installation incidents
• Chevron Ballymore - Rigid jumper and manifold installation analysis with dynamic and resonance analysis
• Talos Venice Limerock Field - Static umbilical installation analysis in 1,050m water depth
• Circular BOP Design - Detailed FEA analysis for seals and connectors using nonlinear elastic-plastic analysis
• 42-inch Pipeline Installation (Venezuela) - FEA analysis for thin-walled pipe (D/t=67) outside DNV code regime
• Digital Twin Development - Automated analysis for drilling risers, TTRs, SCRs, LWCR with fitness for service per API 579:2016
• Trendsetter Intervention System - Design for 1,500-10,000 ft water depths and 15,000-17,500 psi pressures
• Corrosion Simulation - COMSOL modeling of damaged subsea mooring lines with electrochemical analysis
Client References: ExxonMobil, Chevron, Talos Energy

Sep 2017 - Dec 2020: Engineering SME for Data Science Team, Occidental Petroleum, Houston, TX
• Production Optimization Platform - Deployed 50+ physics-based algorithms monitoring 20,000 wells in real-time (99.8% uptime)
• Real-Time Drilling Analytics - Implemented bit position estimation (±2ft accuracy at 10,000ft), reduced drilling time by 18%
• Production Surveillance System - Automated dynacard analysis, reduced ESP failures by 25%, improved well test validation from 70% to 95%
• Deployed edge computing solution processing 1TB daily with <100ms latency
• Achieved 35% reduction in non-productive time, saving $15M+ annually
Client Reference: Occidental Petroleum

Aug 2003 - Jun 2015: Engineering Lead, 2H Offshore Inc, Houston, TX
• BP Macondo Emergency Response - Engineering Manager delivering complete containment riser design in 8 weeks (vs. typical 3 years), system contained 15,000 bpd during 2010 Gulf oil spill
• Chevron Bangka SCR/Flexible Riser Design ($45M FEED) - Led design for 8 production risers in 300ft water depth, 18% weight reduction ($3M savings)
• Major Project Portfolio (100+ projects): BP Thunder Horse/Atlantis/Horn Mountain integrity management, Chevron Jack/St.Malo SCRs, ENI Devils Tower TTRs, Shell/BP HPHT designs
• API Standards - Committee member and author for API-RP-16Q (2016 release)
• Global Clients: BP, Shell, Chevron, ExxonMobil, ENI, Murphy, Reliance across GoM, West Africa, Asia-Pacific
Client References: BP, Shell, Chevron, ExxonMobil

Awards & Recognition:
• Multiple accolades for customer service and satisfaction
• 100% track record for on-time and within-budget project delivery""",
        
        'languages': """English - Excellent (speaking, reading, writing)
Telugu - Excellent (speaking, reading, writing)
Hindi - Good (speaking, reading, writing)""",
        
        'date': datetime.now().strftime('%d/%B/%Y'),
        'full_name': 'Vamsee Achanta, P.E.'
    }
    
    # Map fields to paragraphs by searching for keywords
    field_mapping = {
        'Name of Consultant:': 'name',
        'Profession:': 'profession',
        'Date of Birth:': 'dob',
        'Nationality:': 'nationality',
        'Membership in Professional Societies:': 'memberships',
        'Function assigned in Consultant': 'function'
    }
    
    # Populate fields
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Handle simple field replacements
        for keyword, data_key in field_mapping.items():
            if keyword in text and data_key in cv_data:
                # Replace the paragraph text
                para.text = f"{keyword}\t{cv_data[data_key]}"
        
        # Handle multi-paragraph sections
        if 'Key Qualifications:' in text:
            # Insert key qualifications after the next paragraph
            if i + 2 < len(doc.paragraphs):
                doc.paragraphs[i + 2].text = cv_data['key_qualifications']
        
        elif 'Education:' in text:
            # Insert education after the next paragraph
            if i + 2 < len(doc.paragraphs):
                doc.paragraphs[i + 2].text = cv_data['education']
        
        elif 'Employment Record:' in text:
            # Insert employment after the next paragraph
            if i + 2 < len(doc.paragraphs):
                doc.paragraphs[i + 2].text = cv_data['employment']
        
        elif 'Languages:' in text:
            # Insert languages after the next paragraph
            if i + 2 < len(doc.paragraphs):
                doc.paragraphs[i + 2].text = cv_data['languages']
        
        elif 'Day/Month/Year' in text:
            para.text = para.text.replace('Day/Month/Year', cv_data['date'])
        
        elif 'Full name of Consultant:' in text:
            para.text = f"Full name of Consultant: {cv_data['full_name']}"
    
    # Save the populated document
    output_path = 'cv/acma/02_populated/Vamsee_Achanta_CV_ACMA.docx'
    doc.save(output_path)
    print(f"CV populated successfully: {output_path}")
    return output_path

if __name__ == '__main__':
    populate_cv_template()